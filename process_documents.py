from __future__ import annotations

import json
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable

from docx import Document
from pypdf import PdfReader

BASE_DIR = Path(__file__).resolve().parent.parent
OUTPUT_FILE = BASE_DIR / "dados_processados" / "chunks.jsonl"

RESPONSAVEIS = {
    "01_RH": "Mariana Oliveira - Coordenação de RH",
    "02_Operacional": "Rafael Mendes - Coordenação de Operações",
    "03_Financeiro": "Juliana Costa - Coordenação Financeira",
    "04_Cursos_e_Certificados": "Camila Ribeiro - Coordenação Acadêmica",
    "05_Comercial": "Lucas Almeida - Coordenação Comercial",
}

CATEGORIAS = {
    "01_RH": "RH",
    "02_Operacional": "Operacional",
    "03_Financeiro": "Financeiro",
    "04_Cursos_e_Certificados": "Cursos e Certificados",
    "05_Comercial": "Comercial",
}

ARQUIVO_CATEGORIA = {
    "01_Regimento_do_Estudante.docx": ("RH", "Mariana Oliveira - Coordenação de RH"),
    "02_Politica_de_Reembolso_de_Matriculas.docx": ("Financeiro", "Juliana Costa - Coordenação Financeira"),
    "03_FAQ_Cursos_e_Certificados.docx": ("Cursos e Certificados", "Camila Ribeiro - Coordenação Acadêmica"),
    "04_Guia_de_Uso_da_Plataforma.docx": ("Operacional", "Rafael Mendes - Coordenação de Operações"),
    "05_Programa_de_Bolsas_e_Afiliados.docx": ("Comercial", "Lucas Almeida - Coordenação Comercial"),
}

CHUNK_SIZE = 900
CHUNK_OVERLAP = 120


def clean_text(text: str) -> str:
    """Remove ruídos básicos sem alterar o conteúdo semântico."""
    text = text.replace("\u00a0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def split_text(text: str, size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    """Divide preferencialmente por parágrafos/frases, com pequena sobreposição."""
    if not text:
        return []
    if len(text) <= size:
        return [text]

    paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
    chunks: list[str] = []
    current = ""

    for paragraph in paragraphs:
        candidate = f"{current}\n{paragraph}".strip() if current else paragraph
        if len(candidate) <= size:
            current = candidate
            continue

        if current:
            chunks.append(current)

        # Evita perder contexto quando um parágrafo sozinho é grande.
        if len(paragraph) > size:
            start = 0
            while start < len(paragraph):
                end = min(start + size, len(paragraph))
                piece = paragraph[start:end].strip()
                if piece:
                    chunks.append(piece)
                if end >= len(paragraph):
                    break
                start = max(end - overlap, start + 1)
            current = ""
        else:
            previous = chunks[-1] if chunks else ""
            context = previous[-overlap:] if previous else ""
            current = f"{context}\n{paragraph}".strip() if context else paragraph

    if current:
        chunks.append(current)

    return chunks


def metadata_for(path: Path) -> dict:
    # Categorias/responsáveis vêm do mapeamento organizacional da etapa 2.
    category_key = path.parent.name if path.parent.name in CATEGORIAS else None
    categoria = CATEGORIAS.get(category_key)
    responsavel = RESPONSAVEIS.get(category_key)

    # Fallback para os arquivos atuais, caso ainda estejam na pasta principal.
    if not categoria or not responsavel:
        mapped = ARQUIVO_CATEGORIA.get(path.name)
        if mapped:
            categoria, responsavel = mapped

    stat = path.stat()
    updated = datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc).isoformat()
    return {
        "arquivo_original": path.name,
        "formato": path.suffix.lower().lstrip("."),
        "categoria": categoria or "Não classificado",
        "responsavel": responsavel or "Não definido",
        "ultima_modificacao_arquivo": updated,
        "fonte": "pasta local da Ateliê Forma Online",
    }


def extract_docx(path: Path) -> list[dict]:
    doc = Document(path)
    items: list[dict] = []

    for idx, paragraph in enumerate(doc.paragraphs, start=1):
        text = clean_text(paragraph.text)
        if text:
            items.append({
                "texto": text,
                "localizacao": f"parágrafo {idx}",
            })

    for table_idx, table in enumerate(doc.tables, start=1):
        rows = []
        for row in table.rows:
            cells = [clean_text(cell.text) for cell in row.cells]
            rows.append(" | ".join(cells))
        table_text = "\n".join(r for r in rows if r)
        if table_text:
            items.append({
                "texto": clean_text(table_text),
                "localizacao": f"tabela {table_idx}",
            })

    return items


def extract_pdf(path: Path) -> list[dict]:
    reader = PdfReader(str(path))
    items: list[dict] = []
    for page_number, page in enumerate(reader.pages, start=1):
        text = clean_text(page.extract_text() or "")
        if text:
            items.append({
                "texto": text,
                "localizacao": f"página {page_number}",
            })
    return items


def extract_file(path: Path) -> list[dict]:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return extract_docx(path)
    if suffix == ".pdf":
        return extract_pdf(path)
    return []


def iter_documents() -> Iterable[Path]:
    for path in sorted(BASE_DIR.rglob("*")):
        if not path.is_file():
            continue
        if "processamento" in path.parts or "dados_processados" in path.parts:
            continue
        if path.suffix.lower() in {".docx", ".pdf"}:
            yield path


def build_chunks() -> list[dict]:
    generated_at = datetime.now(timezone.utc).isoformat()
    results: list[dict] = []
    chunk_id = 1

    for path in iter_documents():
        metadata = metadata_for(path)
        blocks = extract_file(path)
        for block in blocks:
            for piece in split_text(block["texto"]):
                results.append({
                    "id": f"chunk-{chunk_id:04d}",
                    "texto": piece,
                    "metadados": {
                        **metadata,
                        "localizacao": block["localizacao"],
                        "processado_em": generated_at,
                    },
                })
                chunk_id += 1

    return results


def main() -> None:
    chunks = build_chunks()
    OUTPUT_FILE.parent.mkdir(parents=True, exist_ok=True)
    with OUTPUT_FILE.open("w", encoding="utf-8") as f:
        for chunk in chunks:
            f.write(json.dumps(chunk, ensure_ascii=False) + "\n")
    print(f"Documentos processados: {len(list(iter_documents()))}")
    print(f"Chunks gerados: {len(chunks)}")
    print(f"Arquivo: {OUTPUT_FILE}")


if __name__ == "__main__":
    main()
